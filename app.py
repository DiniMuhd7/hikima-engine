import os
import torch
import random
import docx
import PyPDF2
import re
import glob
import json
import tempfile
import math
from torch import nn
from torch.nn import functional as F
from torch.utils.data import DataLoader
import numpy as np
import commons
import utils
import argparse
from data_utils import TextAudioLoader, TextAudioCollate, TextAudioSpeakerLoader, TextAudioSpeakerCollate
from models import SynthesizerTrn
from scipy.io.wavfile import write
import scipy.io.wavfile as wavsaver
import argparse
import uuid
from pydub import AudioSegment
from moviepy.editor import VideoFileClip, AudioFileClip
import whisper
import spacy
from spacy_syllables import SpacySyllables
from tqdm import tqdm
from translate import Translator
import tempfile
import subprocess
import locale
import ast
import requests
from retry import retry
from pytube import YouTube
from dotenv import load_dotenv
from urllib.parse import urlparse, parse_qs
from sqlalchemy import create_engine, Column, Integer, Float, String, Text
from sqlalchemy.orm import sessionmaker
from sqlalchemy.exc import NoResultFound
from sqlalchemy.ext.declarative import declarative_base
locale.getpreferredencoding = lambda: "UTF-8"

from flask import (Flask, redirect, render_template, request, session,
                   send_from_directory, url_for, jsonify, make_response)

from flask_oauthlib.client import OAuth

#from InferenceInterfaces.ToucanTTSInterface import ToucanTTSInterface

# Load environment variables from the .env file in the current directory
load_dotenv()

# configs
FLUTTERWAVE_API_LIVE_SECRET_KEY = 'FLWSECK-ee918e0f29b93902dc17bef1abda7d42-18ba9552b35vt-X'
FLUTTERWAVE_API_WEBHOOK_SIGNATURE = 'ee918e0f29b906efcfbbee16'
FLUTTERWAVE_API_REDIRECT_URL = 'https://owltape.azurewebsites.net'
APP_SECRET_KEY = '@\xd6\x01FC#\xf98E\xe0\xff\xb9\xc8\xf4z\xad\x06\xb2"\xfa\xc0\x91\x94\xce'
DATABASE_MAIN_URI = 'sqlite:///hikima.db'
GOOGLE_CONSUMER_KEY = '815941416085-1lgmbst4a43gn86pufg3ces27t6qbrdi.apps.googleusercontent.com'
GOOGLE_CONSUMER_SECRET = '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'

app = Flask(__name__)
app.secret_key = APP_SECRET_KEY  # Set a secret key for session encryption.
app.config['SQLALCHEMY_DATABASE_URI'] = DATABASE_MAIN_URI  # Replace with your actual database URL

# Create db instance for handling database
engine = create_engine(DATABASE_MAIN_URI, echo=True)
# Define a base class for declarative class definitions
Base = declarative_base()

firmpay_code = "/"+str(uuid.uuid4())

# Define a SQLAlchemy model for payment records
class BalanceRecord(Base):
    __tablename__ = 'balance_record'  # Set the table name

    id = Column(Integer, primary_key=True)
    user_name = Column(String(255), nullable=False)
    user_email = Column(String(255), nullable=False)
    tx_ref = Column(String(255), nullable=False)
    amount = Column(Float, nullable=False)
    tamount = Column(Float, nullable=False)
    payment_status = Column(String(50), nullable=False)

    def __init__(self, user_name, user_email, tx_ref, amount, tamount, payment_status):
        self.user_name = user_name
        self.user_email = user_email
        self.tx_ref = tx_ref
        self.amount = amount
        self.tamount = tamount
        self.payment_status = payment_status


class CrowdsourceRecord(Base):
    __tablename__ = 'crowdsource_entity'  # table name

    id = Column(Integer, primary_key=True)
    user_email = Column(String(255), nullable=False)
    media_type = Column(String(50), nullable=False)
    media_path = Column(String(255), nullable=False)
    audio_path = Column(String(255), nullable=False)
    media_transcript = Column(Text, nullable=False)
    media_lang = Column(String(100), nullable=False)
    audio_status = Column(String(50), nullable=False)
    translate_from = Column(String(50), nullable=False)
    translate_to = Column(String(50), nullable=False)
    media_words = Column(Integer, nullable=False)
    speech_engine = Column(String(50), nullable=False)
    translation_fee = Column(Float, nullable=False)
    media_translator = Column(String(255), nullable=False)
    crowdsource_status = Column(String(50), nullable=False)

    def __init__(self, user_email, media_type, media_path, audio_path, media_transcript, media_lang, audio_status, translate_from, translate_to, media_words, speech_engine, translation_fee, media_translator, crowdsource_status):
        self.user_email = user_email
        self.media_type = media_type
        self.media_path = media_path
        self.audio_path = audio_path
        self.media_transcript = media_transcript
        self.media_lang = media_lang
        self.audio_status = audio_status
        self.translate_from = translate_from
        self.translate_to = translate_to
        self.media_words = media_words
        self.speech_engine = speech_engine
        self.translation_fee = translation_fee
        self.media_translator = media_translator
        self.crowdsource_status = crowdsource_status


# Google Sign in Auth Credentials
oauth = OAuth()
google = oauth.remote_app(
    "google",
    consumer_key=GOOGLE_CONSUMER_KEY,
    consumer_secret=GOOGLE_CONSUMER_SECRET,
    request_token_params={
        "scope": "openid profile email"
    },
    base_url="https://www.googleapis.com/oauth2/v1/",
    request_token_url=None,
    access_token_method="POST",
    access_token_url="https://accounts.google.com/o/oauth2/token",
    authorize_url="https://accounts.google.com/o/oauth2/auth",
)

# NLTK Spacy models declaration
spacy_models = {
    "english": "en_core_web_sm",
    "german": "de_core_news_sm",
    "french": "fr_core_news_sm",
    "italian": "it_core_news_sm",
    "catalan": "ca_core_news_sm",
    "chinese": "zh_core_web_sm",
    "croatian": "hr_core_news_sm",
    "danish": "da_core_news_sm",
    "dutch": "nl_core_news_sm",
    "finnish": "fi_core_news_sm",
    "greek": "el_core_news_sm",
    "japanese": "ja_core_news_sm",
    "korean": "ko_core_news_sm",
    "lithuanian": "lt_core_news_sm",
    "macedonian": "mk_core_news_sm",
    "polish": "pl_core_news_sm",
    "portuguese": "pt_core_news_sm",
    "romanian": "ro_core_news_sm",
    "russian": "ru_core_news_sm",
    "spanish": "es_core_news_sm",
    "slovenian": "sl_core_news_sm",
    "swedish": "sv_core_news_sm",
    "ukrainian": "uk_core_news_sm"
}


@app.route('/')
def index():
    if 'google_token' in session:
        #return f"Logged in as {session['username']} | <a href='/logout'>Logout</a>"
        return redirect(url_for('speech'))
    else:
        #return "Not logged in | <a href='/login'>Login</a>"
        return redirect(url_for('login'))



@app.route('/welcome')
def welcome():
    return render_template('index.html')


"""
@app.route('/vocal')
def vocal():
    sentences = [
        "Bringing technology to your doorstep in an accelarating time.",
        "Under whose authority does a machine learning dataset is built?",
        "Overfitting is a major problem in training an AI model, in what ways can this be addressed?"
    ]
 
    return render_template('profile.html',sentences=sentences)


@app.route('/record', methods=['POST'])
def record():
    user_id = request.form.get('user_id')
    sentence = request.form.get('sentence')

    # Replace this with your implementation to handle the recording and saving logic
    # Save the recording to a directory named after the user ID
    audio_file = request.files['audio']
    if audio_file:
        directory = f"recordings/{user_id}"
        os.makedirs(directory, exist_ok=True)
        filepath = os.path.join(directory, f"{sentence.replace(' ', '_')}.wav")
        audio_file.save(filepath)

        return jsonify(success=True, message="Recording saved successfully!")

    return jsonify(success=False, message="Failed to save the recording.")
"""

@app.route('/logout')
def logout():
    session.pop('google_token', None)
    session.pop("user_fullname", None)
    session.pop("user_email", None)
    session.pop("user_id", None)
    
    return redirect('/login')


@app.route('/favicon.ico')
def favicon():
    return send_from_directory(os.path.join(app.root_path, 'static'),
                               'favicon.ico', mimetype='image/vnd.microsoft.icon')


@app.route('/login', methods=['GET', 'POST'])
def login():
    if 'google_token' in session:
        #return f"Logged in as {session['username']} | <a href='/logout'>Logout</a>"
        return redirect(url_for('speech'))
    return render_template('login.html', message='')


@app.route('/login/auth', methods=['GET', 'POST'])
def google_auth_request():
    return google.authorize(callback=url_for("authorized", _external=True))

@google.tokengetter
def get_google_oauth_token():
    return session.get("google_token")

@app.route("/login/authorized", methods=['GET', 'POST'])
def authorized():
    response = google.authorized_response()
    if response is None or response.get("access_token") is None:
        return "Access denied: reason={} error={}".format(
            request.args["error_reason"], request.args["error_description"]
        )
    session["google_token"] = (response["access_token"], "")
    me = google.get("userinfo")

    user_fullname = me.data.get("name")
    user_email = me.data.get("email")
    user_id = me.data.get("id")

    session["user_fullname"] = user_fullname
    session["user_email"] = user_email
    session["user_id"] = user_id

    try:
        # Create a session to interact with the database
        Session = sessionmaker(bind=engine)
        db_session = Session()

        user_name = session["user_fullname"]
        user_email = session["user_email"]
        
        records = db_session.query(BalanceRecord.user_email).filter(BalanceRecord.user_email == user_email).first()
        # Print the retrieved records
        if not records:
            # No record found, add user's account with the purchased amount
            balance_record = BalanceRecord(
                user_name=user_name,
                user_email=user_email,
                tx_ref='unknown',
                amount=0.0,
                tamount=0.0,
                payment_status='insufficient'
            )
            db_session.add(balance_record)
            db_session.commit()
            print("No records found in the database. A new user will be added.")
            return redirect(url_for('speech')) # Redirect to speech URL 
        else:
            return redirect(url_for('speech')) # Redirect to speech URL 
    except NoResultFound:
        print("No records found in the database.")
    except Exception as e:
        print(f"An error occurred: {e}")
    finally:
        db_session.close()


# Read texts
def read_texts(model_id, sentence, filename, device="cpu", language="", speaker_reference=None, faster_vocoder=False):
    tts = ToucanTTSInterface(device=device, tts_model_path=model_id, faster_vocoder=faster_vocoder)
    tts.set_language(language)
    if speaker_reference is not None:
        tts.set_utterance_embedding(speaker_reference)
    if type(sentence) == str:
        sentence = [sentence]
    tts.read_to_file(text_list=sentence, file_location=filename)
    del tts


# Setup model for inference experiment
def the_raven(version, model_id, exsentence, exec_device="cpu", speed_over_quality=False, speaker_reference=None, langcode=""):
    if langcode != "" and langcode == "ha":
        twicklang = "sw"
        read_texts(model_id="Hausa",
               sentence=exsentence,
               filename=f"static/audios/madugu_{version}.wav",
               device=exec_device,
               language=twicklang,
               speaker_reference=speaker_reference,
               faster_vocoder=speed_over_quality)
        #return render_template('inference.html', audio_visible=True, audio_file=audio_path)
    elif langcode != "" and langcode == "sw":
        read_texts(model_id="Swahili",
               sentence=exsentence,
               filename=f"static/audios/tai_{version}.wav",
               device=exec_device,
               language=langcode,
               speaker_reference=speaker_reference,
               faster_vocoder=speed_over_quality)


# Initiate inference for experiment
@app.route('/speech', methods=['GET', 'POST'])
def speech():
    sentences = []
    if 'google_token' in session:
        #return f"Logged in as {session['username']} | <a href='/logout'>Logout</a>"
        #redirect(url_for('inference'))
        os.makedirs("static/audios", exist_ok=True)

        user_fullname = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        tamount = show_credit_balance()
        greetings = f"Hello, {user_fullname}"

        sentence = request.form.get('syn')
        langtag = request.form.get('tag')
        langiso = request.form.get('isocode')
        #print(str(langtag))
    
        digit = random.randint(0, 1000)
        exec_device = "cuda" if torch.cuda.is_available() else "cpu"
        message = "Your audio is ready. Listen to Speech."
        #messagex = f"Running on {exec_device}. Synthetizing..."

        try:
            # Create a session to interact with the database
            Session = sessionmaker(bind=engine)
            db_session = Session()
        
            # Create a session to interact with the database
            record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
            # Print the retrieved records
            if not record:
                # No record found, add user's account with the purchased amount
                print("No record found in the database.")
                return redirect(url_for('speech')) # Redirect to speech URL 
            else:
                # Retrieve user's total amount
                user_balance = record.tamount
                #db_session.commit()
                if user_balance is not None and user_balance <= 5.0:
                    error_message = "Sorry "+user_fullname+"!"+" Your current balance is insufficient. Purchase more credit units."
                    print("Sorry "+user_fullname+"!"+" Your credit balance is insufficient. Purchase more credit units.")
                    #return render_template('inference.html', greetings=greetings, balance=tamount, error_visible=True, error_message=error_message)
                else:
                    if sentence is not None and langtag is not None and langiso is not None:
                        #if langtag != "" and langtag == "ha":
                        sentences.clear() 
                        sentences.append(sentence)

                        audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                        create_speech_from_text(sentences, langtag, langiso, audio_path)

                        file_size = get_file_size(audio_path) # Get the file size in bytes
                        if file_size:
                            print(f"File size: {file_size} bytes. File size is valid.")
                            task_cost_per_mb = 1.5  # Cost per MB in credits
                            file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                            total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost
                                
                            # Modify and substract user's total amount for usage cost
                            current_balance = user_balance - total_task_cost
                            current_balance = round(current_balance, 2)
                            record.tamount = current_balance # Update user balance in database
                            db_session.commit()

                        #the_raven(version=digit,
                        #    model_id="Hausa",
                        #    exec_device=exec_device,
                        #    speed_over_quality=exec_device != "cuda",
                        #    exsentence=sentence,
                        #    langcode=langtag)
                        #audio_path = os.path.join('static/audios',f'madugu_{digit}.wav')
                        return render_template('inference.html',  greetings=greetings, balance=tamount, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                        #elif langtag != "" and langtag == "sw":
                        #sentences.clear() 
                        #sentences.append(sentence)

                        #audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                        #create_speech_from_text(sentences, langtag, langiso, audio_path)
                        #the_raven(version=digit,
                        #    model_id="Swahili",
                        #    exec_device=exec_device,
                        #    speed_over_quality=exec_device != "cuda",
               	        #    exsentence=sentence,
                        #    langcode=langtag)
                        #audio_path = os.path.join('static/audios',f'tai_{digit}.wav')
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            db_session.rollback()
            print(f"An error occurred: {e}")
        finally:
            db_session.close()

        #return render_template('inference.html', greetings=greetings, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
        return render_template('inference.html', greetings=greetings, balance=tamount, langtag = langtag, audio_visible=False, audio_file='', message='')
    else:
        #return "Not logged in | <a href='/login'>Login</a>"
        return redirect(url_for('login'))


def create_table():
    with app.app_context():    
        db_session = scoped_session(sessionmaker(autocommit=False, autoflush=False, bind=db.engine))
        # Payment was successful, update user's account with the purchased amount
        balance_record = BalanceRecord(
            user_name="",
            user_email="",
            tx_ref="",
            amount=0.00,
            tamount=0.00,
            payment_status=""
        )

        db_session.add(balance_record)
        db_session.commit()


def show_credit_balance():
    if 'google_token' in session:
        user_fullname = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")
        # Perform additional checks and validations if needed
        
        # Create a session to interact with the database
        Session = sessionmaker(bind=engine)
        db_session = Session()
        # Payment was successful, update user's account with the purchased amount
        try:
            records = db_session.query(BalanceRecord.tamount).filter(BalanceRecord.user_email == user_email).first()
            db_session.commit()
            # Print the retrieved records
            if not records:
                balance = 0.00
                print("No records found in the database.")
                return balance
            else:
                for record in records:
                    return record
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            print(f"An error occurred: {e}")
        finally:
            db_session.close()
    else:
        return redirect(url_for('login'))
    
# File extensions
def allowed_file(filename):
    # Check if the file extension is allowed
    allowed_extensions = {'txt', 'pdf', 'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


# File extensions
def allowed_media_file(filename):
    # Check if the media file extension is allowed
    allowed_extensions = {'mp4', 'mp3'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def check_file_extension(file_path, extension):
    file_ext = os.path.splitext(file_path)[1]
    return file_ext.lower() == extension.lower()


# Read files
txt_archive = []
def read_txtfile(file_path):
    txt_archive.clear()
    with open(file_path, 'r') as file:
        for line in file:
            txt_archive.append(line.strip())
    return txt_archive

pdf_archive = []
def read_pdffile(file_path):
    pdf_archive.clear()
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfFileReader(file)
        for page_number in range(pdf_reader.numPages):
            page = pdf_reader.getPage(page_number)
            text = page.extract_text()
            pdf_archive.append(text)

    return pdf_archive

doc_archive = []
def read_docfile(file_path):
    doc_archive.clear()
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        doc_archive.append(text)
    
    return doc_archive

docx_archive = []
def read_docxfile(file_path):
    docx_archive.clear()
    doc = docx.Document(file_path)
    for paragraph in doc.paragraphs:
        text = paragraph.text
        docx_archive.append(text)
    
    return docx_archive


def get_file_size(file_path):
    try:
        file_size = os.path.getsize(file_path)
        return file_size
    except Exception as e:
        print(f"Error getting file size: {e}")
        return None

def download_youtube_video(video_url, save_path):
    try:
        # Create a YouTube object
        yt = YouTube(video_url)
        
        # Get the highest resolution stream (you can customize this)
        stream = yt.streams.get_highest_resolution()
        
        # Download the video
        filename = stream.download(output_path=save_path)
        
        print("Download successful!")
        return filename
    except Exception as e:
        print("An error occurred:", str(e))
        return "Check your network connection and try again"

def check_currency(currency_code, amount):

    # Fetch the latest exchange rates
    response = requests.get('https://www.floatrates.com/daily/usd.json')
    data = response.json()

    current_amount = 0
    # Check if the API call was successful
    if currency_code == "usd":
        current_amount = amount
    elif currency_code in data:
        currency_rate = data[currency_code]['rate']
        exchange_rate = 1 / currency_rate
        usd_equivalent = amount / exchange_rate
        current_amount = usd_equivalent
        print(f"The equivalent amount in {currency_code} is: {usd_equivalent:.2f}")
    else:
        print(f"Exchange rate for {currency_code} not available.")

    return current_amount


@retry(tries=3, delay=2, backoff=2)
@app.route('/slatepay', methods=['GET', 'POST'])
def slatepay():
    if 'google_token' in session:
        user_name = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        amount_in_usd = 4.99
        currency = request.form.get('currency')
        current_amount = check_currency(currency.lower(), amount_in_usd)

        payload = {
            "tx_ref": str(uuid.uuid4()),
            "amount": current_amount,  # Amount in USD
            "currency": currency,
            "redirect_url": FLUTTERWAVE_API_REDIRECT_URL+""+firmpay_code,
            "customer": {
                "email": user_email,
                "phonenumber": "08100000000",
                "name": user_name
            },
            "meta": {
                "user_id": user_id
            },
            "customizations": {
                "title": "Hikima AI",
                "logo": "https://pbs.twimg.com/profile_images/1671146308294287361/9Q8K_iwk_400x400.jpg"
            }
        }

        headers = {
            "Authorization": f"Bearer {FLUTTERWAVE_API_LIVE_SECRET_KEY}",
            "Content-Type": "application/json"
        }

        try:
            # Create a session to interact with the database
            Session = sessionmaker(bind=engine)
            db_session = Session()
        
            # Create a session to interact with the database
            record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
            # Print the retrieved records
            if not record:
                # No record found, add user's account with the purchased amount
                print("No record found in the database. A new user will be added.")
                return redirect(url_for('purchase')) # Redirect to speech URL 
            else:
                # Update user's record
                record.amount = amount_in_usd
                db_session.commit()
                print("User purchase amount updated successfully")
                
                # Initiate payment gateway
                response = requests.post("https://api.flutterwave.com/v3/payments", json=payload, headers=headers)

                data = response.json()
                print("Payment initiation response:", data)

                if data.get("status") == "success":
                    payment_link = data["data"]["link"]
                    print("Payment Link:", payment_link)
                    return redirect(payment_link)
                else:
                    # Handle error case here
                    print("Payment initiation failed.")
                    return redirect(url_for('purchase')) # Redirect to speech URL 
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            db_session.rollback()
            print(f"An error occurred: {e}")
        finally:
            db_session.close()     
    else:
        return redirect(url_for('login'))



@retry(tries=3, delay=2, backoff=2)
@app.route('/goldpay', methods=['GET', 'POST'])
def goldpay():
    if 'google_token' in session:
        user_name = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        amount_in_usd = 9.99
        currency = request.form.get('currency')
        current_amount = check_currency(currency.lower(), amount_in_usd)

        payload = {
            "tx_ref": str(uuid.uuid4()),
            "amount": current_amount,  # Amount in USD
            "currency": currency,
            "redirect_url": FLUTTERWAVE_API_REDIRECT_URL+""+firmpay_code,
            "customer": {
                "email": user_email,
                "phonenumber": "08100000000",
                "name": user_name
            },
            "meta": {
                "user_id": user_id
            },
            "customizations": {
                "title": "Hikima AI",
                "logo": "https://pbs.twimg.com/profile_images/1671146308294287361/9Q8K_iwk_400x400.jpg"
            }
        }

        headers = {
            "Authorization": f"Bearer {FLUTTERWAVE_API_LIVE_SECRET_KEY}",
            "Content-Type": "application/json"
        }

        try:
            # Create a session to interact with the database
            Session = sessionmaker(bind=engine)
            db_session = Session()
        
            # Create a session to interact with the database
            record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
            # Print the retrieved records
            if not record:
                # No record found, add user's account with the purchased amount
                print("No record found in the database. A new user will be added.")
                return redirect(url_for('purchase')) # Redirect to speech URL 
            else:
                # Update user's record
                record.amount = amount_in_usd
                db_session.commit()
                print("User purchase amount updated successfully")
                
                # Initiate payment gateway
                response = requests.post("https://api.flutterwave.com/v3/payments", json=payload, headers=headers)

                data = response.json()
                print("Payment initiation response:", data)

                if data.get("status") == "success":
                    payment_link = data["data"]["link"]
                    print("Payment Link:", payment_link)
                    return redirect(payment_link)
                else:
                    # Handle error case here
                    print("Payment initiation failed.")
                    return redirect(url_for('purchase')) # Redirect to speech URL 
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            db_session.rollback()
            print(f"An error occurred: {e}")
        finally:
            db_session.close()     
    else:
        return redirect(url_for('login'))



@retry(tries=3, delay=2, backoff=2)
@app.route('/pilotpay', methods=['GET', 'POST'])
def pilotpay():
    if 'google_token' in session:
        user_name = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        amount_in_usd = 1.00
        currency = request.form.get('currency')
        current_amount = check_currency(currency.lower(), amount_in_usd)

        payload = {
            "tx_ref": str(uuid.uuid4()),
            "amount": current_amount,  # Amount in USD
            "currency": currency,
            "redirect_url": FLUTTERWAVE_API_REDIRECT_URL+""+firmpay_code,
            "customer": {
                "email": user_email,
                "phonenumber": "08100000000",
                "name": user_name
            },
            "meta": {
                "user_id": user_id
            },
            "customizations": {
                "title": "Hikima AI",
                "logo": "https://pbs.twimg.com/profile_images/1671146308294287361/9Q8K_iwk_400x400.jpg"
            }
        }

        headers = {
            "Authorization": f"Bearer {FLUTTERWAVE_API_LIVE_SECRET_KEY}",
            "Content-Type": "application/json"
        }

        try:
            # Create a session to interact with the database
            Session = sessionmaker(bind=engine)
            db_session = Session()
        
            # Create a session to interact with the database
            record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
            # Print the retrieved records
            if not record:
                # No record found, add user's account with the purchased amount
                print("No record found in the database. A new user will be added.")
                return redirect(url_for('purchase')) # Redirect to speech URL 
            else:
                # Update user's record
                record.amount = amount_in_usd
                db_session.commit()
                print("User purchase amount updated successfully")
                
                # Initiate payment gateway
                response = requests.post("https://api.flutterwave.com/v3/payments", json=payload, headers=headers)

                data = response.json()
                print("Payment initiation response:", data)

                if data.get("status") == "success":
                    payment_link = data["data"]["link"]
                    print("Payment Link:", payment_link)
                    return redirect(payment_link)
                else:
                    # Handle error case here
                    print("Payment initiation failed.")
                    return redirect(url_for('purchase')) # Redirect to speech URL 
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            db_session.rollback()
            print(f"An error occurred: {e}")
        finally:
            db_session.close()     
    else:
        return redirect(url_for('login'))



@app.route(firmpay_code, methods=['GET', 'POST'])
def firmpay():
    try:
         # Create a session to interact with the database
        Session = sessionmaker(bind=engine)
        db_session = Session()
        if 'google_token' in session:
            user_name = session.get("user_fullname")
            user_email = session.get("user_email")
            user_id = session.get("user_id")

            # Extract parameters from the query string
            status = request.args.get('status')
            tx_ref = request.args.get('tx_ref')
            transaction_id = request.args.get('transaction_id')
        
            if status and tx_ref and transaction_id:
                # You have the parameters here, you can now use them as needed
                if status == 'successful':
                    # Create a session to interact with the database
                    record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
                    # Print the retrieved records
                    if not record:
                        # No record found, add user's account with the purchased amount
                        print("No record found in the database.")
                        return redirect(url_for('purchase')) # Redirect to speech URL 
                    else:
                        amount_in_usd = record.amount # Retrive amount in usd
                        if amount_in_usd is not None and amount_in_usd == 1.00:
                            #Update tamount and add credits to user
                            tamount = record.tamount
                            total_amount = tamount + 200
                            record.tamount = total_amount # Update the user's total amount
                            record.tx_ref = tx_ref #Update tx ref
                            record.payment_status = status #Update payment status
                            db_session.commit()
                            print("User credit units added successfully")
                            return redirect(url_for('purchase')) # Redirect to speech URL
                        elif amount_in_usd is not None and amount_in_usd == 4.99:
                            #Update tamount and add credits to user
                            tamount = record.tamount
                            total_amount = tamount + 1200
                            record.tamount = total_amount # Update the user's total amount
                            record.tx_ref = tx_ref #Update tx ref
                            record.payment_status = status #Update payment status
                            db_session.commit()
                            print("User credit units added successfully")
                            return redirect(url_for('purchase')) # Redirect to speech URL
                        elif amount_in_usd is not None and amount_in_usd == 9.99:
                            #Update tamount and add credits to user
                            tamount = record.tamount
                            total_amount = tamount + 3000
                            record.tamount = total_amount # Update the user's total amount
                            record.tx_ref = tx_ref #Update tx ref
                            record.payment_status = status #Update payment status
                            db_session.commit()
                            print("User credit units added successfully")
                            return redirect(url_for('purchase')) # Redirect to speech URL
                        else:
                            print("Somthing went wrong")
                            return 'Somthing went wrong. Try again after 5 minutes.'
                else:
                    return 'Payment not successful. Check your email for more details.' 
            else:
                return redirect(url_for('purchase')) # Redirect to speech URL 
        else:
            return redirect(url_for('login'))
    except NoResultFound:
        print("No records found in the database.")
    except Exception as e:
        db_session.rollback()
        print(f"An error occurred: {e}")
    finally:
        db_session.close()


@app.route('/purchase', methods=['GET', 'POST'])
def purchase():
    if 'google_token' in session:
        tamount = show_credit_balance()
        return render_template('payment.html', balance=tamount)
    else:
        return redirect(url_for('login'))



@app.route('/translate', methods=['GET', 'POST'])
def translate():
    if 'google_token' in session:
        user_fullname = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")
        
        tamount = show_credit_balance()
        greetings = f"Hello, {user_fullname}"
        return render_template('dub.html', greetings=greetings, balance=tamount, error_visible=False, info_visible=False, transdiv_visible=False)
    else:
        return redirect(url_for('login'))


@retry(tries=3, delay=2, backoff=2)
@app.route('/process', methods=['GET', 'POST'])
def process():
    if 'google_token' in session:
        user_fullname = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        tamount = show_credit_balance()
        greetings = f"Hello, {user_fullname}"

        langtag = request.form.get('transtag')
        langiso = request.form.get('isotag')
        langtype = request.form.get('langtype')
        audio_path = request.form.get('path')
        file_root = request.form.get('root')
        audio_status = request.form.get('status')
        translated = request.form.get('translated')
        
        if langtag is None and translated is None and audio_path is None and audio_status is None:
            return
        
        try:
            translated_list = ast.literal_eval(translated)

            transcription = transcribe_audio(audio_path, langtype)
            if transcription is None:
                return
            
            merged_audio, ducked_audio = human_merge_audio(transcription, translated_list, langtype, langtag, langiso, audio_path)
            if merged_audio is None:
                return

            message = "Your media is ready. Not satisfied? Try out (Human Translation) feature."
            serial = str(uuid.uuid4()) 
            if audio_status == "enabled":   
                replace_audio_in_video(file_root, serial, ducked_audio)
            elif audio_status == "disabled":
                replace_audio_in_video(file_root, serial, merged_audio)
            
            output_filename = os.path.splitext(file_root)[0] + serial + "_translated.mp4"  
            # Save the audio file with the same name as the video file but with a ".wav" extension
            #output_filename = os.path.splitext(file_root)[0] + ".wav"
            #save_audio_to_file(merged_audio, output_filename)

            return render_template('dub.html', greetings=greetings, balance=tamount, message=message, video_visible=True, video_file=output_filename)
        except (ValueError, SyntaxError):
            return 'Invalid input format. Please provide a valid list representation.'


@retry(tries=3, delay=2, backoff=2)
@app.route('/translatex', methods=['GET', 'POST'])
def translatex():
    if 'google_token' in session:
        os.makedirs("static/uploads/media", exist_ok=True)

        user_fullname = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        tamount = show_credit_balance()
        greetings = f"Hello, {user_fullname}"

        try:
            # Create a session to interact with the database
            Session = sessionmaker(bind=engine)
            db_session = Session()
        
            # Create a session to interact with the database
            record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
            # Print the retrieved records
            if not record:
                # No record found, add user's account with the purchased amount
                print("No record found in the database.")
                return redirect(url_for('tanslate')) # Redirect to speech URL 
            else:
                # Retrieve user's total amount
                user_balance = record.tamount
                #db_session.commit()
                if user_balance is not None and user_balance <= 1.0:
                    error_message = "Sorry "+user_fullname+"!"+" Your current balance is insufficient."
                    print("Sorry "+user_fullname+"!"+" Your credit balance is insufficient. Purchase more credit units.")
                    return render_template('dub.html', greetings=greetings, balance=tamount, error_visible=True, error_message=error_message)
                else:
                    # Process upload
                    voption = request.form.get('upmethod')
                    if voption == "vmedia":
                        # Check if a file was uploaded
                        if 'media' not in request.files:
                            return 'No file uploaded.'
	
                        file = request.files['media']

                        # Check if the file exists and has an allowed extension
                        if file.filename == '' or not allowed_media_file(file.filename):
                            return 'Invalid file.'
        
                        # Save the uploaded file
                        file_path = os.path.join('static/uploads/media', file.filename)
                        file.save(file_path)

                        max_allowed_size = 50 * 1024 * 1024  # 50 MB in bytes
                        file_size = get_file_size(file_path) # Get the file size in bytes
                        if file_size is not None and file_size <= max_allowed_size:
                            print(f"File size: {file_size} bytes. File size is valid.")
                            #user_balance = 1000  # Get the user's actual balance
                            task_cost_per_mb = 1.5  # Cost per MB in credits
                            file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                            total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost
                            if user_balance >= total_task_cost:
                                print(f"User balance: {user_balance} credits. Task cost: {total_task_cost} credits. Sufficient balance.")
                                vlangtype = request.form.get('vlangtype')
                                tagfrom = request.form.get('tagfrom')
                                tagto = request.form.get('tagto')
                                langiso = request.form.get('isocode')
                                media_status = request.form.get('status')
                                trantype = request.form.get('translate')
                                if vlangtype != "" and tagfrom != "" and tagto != "" and langiso != "" and trantype != "" and media_status != "":
                                    if trantype == "machine":
                                        # Perform machine translation 
                                        #print(f"{vlangtype} {tagfrom} {tagto} {langiso} {trantype}")
                        
                                        # Add tqdm progress bar to show the progress while running task
                                        #progbar = tqdm(total=100, desc="Processing Media File", unit="file")
                        
                                        audio_file = extract_audio_from_video(file_path)
                                        if audio_file is None:
                                            return

                                        transcription = transcribe_audio(audio_file, vlangtype)
                                        if transcription is None:
                                            return

                                        merged_audio, ducked_audio = merge_audio_files(transcription, vlangtype, tagfrom, tagto, langiso, audio_file)
                                        if merged_audio is None:
                                            return
                        
                                        message = "Your media is ready. Not satisfied? Try out (Human Translation) feature."
                                        serial = str(uuid.uuid4()) 
                                        if media_status == "enabled":   
                                            replace_audio_in_video(file_path, serial, ducked_audio)
                                            # Modify and substract user's total amount for usage cost
                                            current_balance = user_balance - total_task_cost
                                            current_balance = round(current_balance, 2)
                                            record.tamount = current_balance # Update user balance in database
                                            db_session.commit()
                                        elif media_status == "disabled":
                                            replace_audio_in_video(file_path, serial, merged_audio)
                                            # Modify and substract user's total amount for usage cost
                                            current_balance = user_balance - total_task_cost
                                            current_balance = round(current_balance, 2)
                                            record.tamount = current_balance # Update user balance in database
                                            db_session.commit()

                                        output_filename = os.path.splitext(file_path)[0] + serial + "_translated.mp4"  
                                        # Save the audio file with the same name as the video file but with a ".wav" extension
                                        #output_filename = os.path.splitext(file_root)[0] + ".wav"
                                        #save_audio_to_file(merged_audio, output_filename)
                        
                                        return render_template('dub.html', greetings=greetings, balance=tamount, message=message, video_visible=True, video_file=output_filename)
                                    elif trantype == "human":
                                        # Perform human translation
                                        #print(f"{vlangtype} {tagfrom} {tagto} {langiso} {trantype}")
                        
                                        lanscript = vlangtype
                                        if lanscript is None:
                                            return

                                        audio_file = extract_audio_from_video(file_path)
                                        if audio_file is None:
                                            return

                                        transcription = transcribe_audio(audio_file, lanscript)
                                        if transcription is None:
                                            return
                        
                                        clean_transcript = get_transcript_chunks(transcription, lanscript)
                                        if clean_transcript is None:
                                            return

                                        # Modify and substract user's total amount for usage cost
                                        current_balance = user_balance - total_task_cost
                                        current_balance = round(current_balance, 2)
                                        record.tamount = current_balance # Update user balance in database
                                        db_session.commit()

                                        return render_template('dub.html', greetings=greetings, balance=tamount, audio_status=media_status, transdiv_visible=True, file_root=file_path, audio_path=audio_file, lang=lanscript, source_text=clean_transcript, target_text=clean_transcript)
                                    elif trantype == "crowdsource":
                                        info_message = "Your media file has been submitted successfully. It is now awaiting translation. You will be notified when the translation is complete."
                                        
                                        lanscript = vlangtype
                                        if lanscript is None:
                                            return

                                        audio_file = extract_audio_from_video(file_path)
                                        if audio_file is None:
                                            return

                                        transcription = transcribe_audio(audio_file, lanscript)
                                        if transcription is None:
                                            return
                        
                                        clean_transcript = get_transcript_chunks(transcription, lanscript)
                                        if clean_transcript is None:
                                            return

                                        # Convert the list to a JSON string
                                        transcript_json = json.dumps(clean_transcript)

                                        # Modify and substract user's total amount for usage cost
                                        current_balance = user_balance - total_task_cost
                                        current_balance = round(current_balance, 2)
                                        record.tamount = current_balance # Update user balance in database
                                        
                                        # Add a new crowdsource translation in the database
                                        crowdsource_record = CrowdsourceRecord(
                                            user_email=user_email,
                                            media_type='.mp4/.mp3',
                                            media_path=file_path,
                                            audio_path=audio_file,
                                            media_transcript=transcript_json,
                                            media_lang=vlangtype,
                                            audio_status=media_status,
                                            translate_from=tagfrom,
                                            translate_to=tagto,
                                            media_words=0,
                                            speech_engine=langiso,
                                            translation_fee=0.05,
                                            media_translator='not assigned',
                                            crowdsource_status='awaiting translation'
                                        )
                                        db_session.add(crowdsource_record)
                                        db_session.commit()
                                                                      
                                        return render_template('dub.html', greetings=greetings, balance=tamount, info_visible=True, info_message=info_message)
                                else:
                                    return 'Please complete all field required'
                            else:
                                error_message = "Sorry "+user_fullname+"!"+" Your current balance is insufficient."
                                print("Sorry "+user_fullname+"!"+" Your credit balance is insufficient. Purchase more credit units.")
                                return render_template('dub.html', greetings=greetings, balance=tamount, error_visible=True, error_message=error_message)
                        else:
                            print(f"File size: {file_size} bytes. File size is too large.")
                            return 'File size is too large'
                    elif voption == "vyoutube":
                        youtube_url = request.form.get('youtube')
                        if youtube_url is None:
                            return 'Youtube video URL required but not provided'
            
                        save_path = "static/uploads/media"
                        yt_filename = download_youtube_video(youtube_url, save_path)

                        # Reference file
                        file_path = os.path.join('static/uploads/media', yt_filename)
                        print(file_path)

                        max_allowed_size = 50 * 1024 * 1024  # 50 MB in bytes
                        file_size = get_file_size(file_path) # Get the file size in bytes
                        if file_size is not None and file_size <= max_allowed_size:
                            print(f"File size: {file_size} bytes. File size is valid.")
                            #user_balance = 1000  # Get the user's actual balance
                            task_cost_per_mb = 1.5  # Cost per MB in credits
                            file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                            total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost
                            if user_balance >= total_task_cost:
                                print(f"User balance: {user_balance} credits. Task cost: {total_task_cost} credits. Sufficient balance.")
                                vlangtype = request.form.get('vlangtype')
                                tagfrom = request.form.get('tagfrom')
                                tagto = request.form.get('tagto')
                                langiso = request.form.get('isocode')
                                media_status = request.form.get('status')
                                trantype = request.form.get('translate')
                                if vlangtype != "" and tagfrom != "" and tagto != "" and langiso != "" and trantype != "" and media_status != "":
                                    if trantype == "machine":
                                        # Perform machine translation 
                                        #print(f"{vlangtype} {tagfrom} {tagto} {langiso} {trantype}")
                        
                                        # Add tqdm progress bar to show the progress while running task
                                        #progbar = tqdm(total=100, desc="Processing Media File", unit="file")
                        
                                        audio_file = extract_audio_from_video(file_path)
                                        if audio_file is None:
                                            return

                                        transcription = transcribe_audio(audio_file, vlangtype)
                                        if transcription is None:
                                            return

                                        merged_audio, ducked_audio = merge_audio_files(transcription, vlangtype, tagfrom, tagto, langiso, audio_file)
                                        if merged_audio is None:
                                            return
                        
                                        message = "Your media is ready. Not satisfied? Try out (Human Translation) feature."
                                        serial = str(uuid.uuid4()) 
                                        if media_status == "enabled":   
                                            replace_audio_in_video(file_path, serial, ducked_audio)
                                            # Modify and substract user's total amount for usage cost
                                            current_balance = user_balance - total_task_cost
                                            current_balance = round(current_balance, 2)
                                            record.tamount = current_balance # Update user balance in database
                                            db_session.commit()
                                        elif media_status == "disabled":
                                            replace_audio_in_video(file_path, serial, merged_audio)
                                            # Modify and substract user's total amount for usage cost
                                            current_balance = user_balance - total_task_cost
                                            current_balance = round(current_balance, 2)
                                            record.tamount = current_balance # Update user balance in database
                                            db_session.commit()

                                        output_filename = os.path.splitext(file_path)[0] + serial + "_translated.mp4"  
                                        # Save the audio file with the same name as the video file but with a ".wav" extension
                                        #output_filename = os.path.splitext(file_root)[0] + ".wav"
                                        #save_audio_to_file(merged_audio, output_filename)
                        
                                        return render_template('dub.html', greetings=greetings, balance=tamount, message=message, video_visible=True, video_file=output_filename)
                                    elif trantype == "human":
                                        # Perform human translation
                                        #print(f"{vlangtype} {tagfrom} {tagto} {langiso} {trantype}")
                        
                                        lanscript = vlangtype
                                        if lanscript is None:
                                            return

                                        audio_file = extract_audio_from_video(file_path)
                                        if audio_file is None:
                                            return

                                        transcription = transcribe_audio(audio_file, lanscript)
                                        if transcription is None:
                                            return
                        
                                        clean_transcript = get_transcript_chunks(transcription, lanscript)
                                        if clean_transcript is None:
                                            return

                                        # Modify and substract user's total amount for usage cost
                                        current_balance = user_balance - total_task_cost
                                        current_balance = round(current_balance, 2)
                                        record.tamount = current_balance # Update user balance in database
                                        db_session.commit()

                                        return render_template('dub.html', greetings=greetings, balance=tamount, audio_status=media_status, transdiv_visible=True, file_root=file_path, audio_path=audio_file, lang=lanscript, source_text=clean_transcript, target_text=clean_transcript)
                                    elif trantype == "crowdsource":
                                        info_message = "Your media file has been submitted successfully. It is now awaiting translation. You will be notified when the translation is complete."
                                        
                                        lanscript = vlangtype
                                        if lanscript is None:
                                            return

                                        audio_file = extract_audio_from_video(file_path)
                                        if audio_file is None:
                                            return

                                        transcription = transcribe_audio(audio_file, lanscript)
                                        if transcription is None:
                                            return
                        
                                        clean_transcript = get_transcript_chunks(transcription, lanscript)
                                        if clean_transcript is None:
                                            return

                                        # Convert the list to a JSON string
                                        transcript_json = json.dumps(clean_transcript)

                                        # Modify and substract user's total amount for usage cost
                                        current_balance = user_balance - total_task_cost
                                        current_balance = round(current_balance, 2)
                                        record.tamount = current_balance # Update user balance in database
                                        
                                        # Add a new crowdsource translation in the database
                                        crowdsource_record = CrowdsourceRecord(
                                            user_email=user_email,
                                            media_type='.mp4/.mp3',
                                            media_path=file_path,
                                            audio_path=audio_file,
                                            media_transcript=transcript_json,
                                            media_lang=vlangtype,
                                            audio_status=media_status,
                                            translate_from=tagfrom,
                                            translate_to=tagto,
                                            media_words=0,
                                            speech_engine=langiso,
                                            translation_fee=0.05,
                                            media_translator='not assigned',
                                            crowdsource_status='awaiting translation'
                                        )
                                        db_session.add(crowdsource_record)
                                        db_session.commit()

                                        return render_template('dub.html', greetings=greetings, balance=tamount, info_visible=True, info_message=info_message)
                                else:
                                    return 'Please complete all field required'
                            else:
                                error_message = "Sorry "+user_fullname+"!"+" Your current balance is insufficient."
                                print("Sorry "+user_fullname+"!"+" Your credit balance is insufficient. Purchase more credit units.")
                                return render_template('dub.html', greetings=greetings, balance=tamount, error_visible=True, error_message=error_message)
                        else:
                            print(f"File size: {file_size} bytes. File size is too large.")
                            return 'File size is too large'
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            db_session.rollback()
            print(f"An error occurred: {e}")
        finally:
            db_session.close()

        return render_template('dub.html', greetings=greetings, balance=tamount)
    else:
        return redirect(url_for('login'))


@retry(tries=3, delay=2, backoff=2)
@app.route('/speechx', methods=['GET', 'POST'])
def speechx():
    if 'google_token' in session:

        user_fullname = session.get("user_fullname")
        user_email = session.get("user_email")
        user_id = session.get("user_id")

        tamount = show_credit_balance()
        greetings = f"Hello, {user_fullname}"

        #return f"Logged in as {session['username']} | <a href='/logout'>Logout</a>"
        #redirect(url_for('inference'))
        os.makedirs("static/audios", exist_ok=True)
        os.makedirs("static/uploads/docs", exist_ok=True)
        #sentence = request.form.get('syn')
        langtag = request.form.get('tagd')
        langiso = request.form.get('isocode')
        #print(str(langtag))

        try:
            # Create a session to interact with the database
            Session = sessionmaker(bind=engine)
            db_session = Session()
        
            # Create a session to interact with the database
            record = db_session.query(BalanceRecord).filter(BalanceRecord.user_email == user_email).first()
            # Print the retrieved records
            if not record:
                # No record found, add user's account with the purchased amount
                print("No record found in the database.")
                return redirect(url_for('speech')) # Redirect to speech URL 
            else:
                # Retrieve user's total amount
                user_balance = record.tamount
                #db_session.commit()
                if user_balance is not None and user_balance <= 5.0:
                    error_message = "Sorry "+user_fullname+"!"+" Your current balance is insufficient."
                    print("Sorry "+user_fullname+"!"+" Your credit balance is insufficient. Purchase more credit units.")
                    return render_template('inference.html', greetings=greetings, balance=tamount, error_visible=True, error_message=error_message)
                else:
                    # Check if a file was uploaded
                    if 'docfile' not in request.files:
                        return 'No file uploaded.'
	
                    file = request.files['docfile']

                    # Check if the file exists and has an allowed extension
                    if file.filename == '' or not allowed_file(file.filename):
                        return 'Invalid file.'

                    # Save the uploaded file
                    file_path = os.path.join('static/uploads/docs', file.filename)
                    file.save(file_path)
    
                    txt_extension='.txt'
                    pdf_extension='.pdf'
                    doc_extension='.doc'
                    docx_extension='.docx'
	
                    is_matching_txt = check_file_extension(file_path, txt_extension)	
                    is_matching_pdf = check_file_extension(file_path, pdf_extension)
                    is_matching_doc = check_file_extension(file_path, doc_extension)
                    is_matching_docx = check_file_extension(file_path, docx_extension)
	
                    audio_path = ""
                    if is_matching_txt is True:
                        #print('Request: '+ str(is_matching_txt)+str(langtag))
                        digit = random.randint(0, 1000)
                        exec_device = "cuda" if torch.cuda.is_available() else "cpu"
                        message = "Your audio is ready. Listen to speech"
            
                        sentences = read_txtfile(file_path)
                        if sentences is not None and langtag is not None and langiso is not None:
                            #if langtag != "" and langtag == "ha":
                            audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            create_speech_from_text(sentences, langtag, langiso, audio_path)
                            
                            file_size = get_file_size(audio_path) # Get the file size in bytes
                            if file_size:
                                print(f"File size: {file_size} bytes. File size is valid.")
                                task_cost_per_mb = 1.5  # Cost per MB in credits
                                file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                                total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost
                                
                                # Modify and substract user's total amount for usage cost
                                current_balance = user_balance - total_task_cost
                                current_balance = round(current_balance, 2)
                                record.tamount = current_balance # Update user balance in database
                                db_session.commit()

                            #the_raven(version=digit,
                            #    model_id="Hausa",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
                            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = send_from_directory(os.path.join(app.root_path, 'audios'),f'madugu_{digit}.wav', mimetype='audio/wav')
                            return render_template('inference.html', greetings=greetings, balance=tamount, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                            #elif langtag != "" and langtag == "sw":
                            #    audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            #    create_speech_from_text(sentences, langtag, langiso, audio_path) 
                            #the_raven(version=digit,
                            #    model_id="Swahili",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
               	            #    exsentence=sentences,
                            #    langcode=langtag)# Check if a file was uploaded
                            #    return render_template('inference.html', greetings=greetings, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                    elif is_matching_pdf is True:
	                    #print('Request: '+ str(is_matching_pdf)+str(langtag))
                        digit = random.randint(0, 1000)
                        exec_device = "cuda" if torch.cuda.is_available() else "cpu"
                        message = "Your audio is ready. Listen to speech"
        
                        sentences = read_pdffile(file_path)
                        if sentences is not None and langtag is not None and langiso is not None:
                            #if langtag != "" and langtag == "ha":
                            audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            create_speech_from_text(sentences, langtag, langiso, audio_path)

                            file_size = get_file_size(audio_path) # Get the file size in bytes
                            if file_size:
                                print(f"File size: {file_size} bytes. File size is valid.")
                                task_cost_per_mb = 1.5  # Cost per MB in credits
                                file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                                total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost
                                
                                # Modify and substract user's total amount for usage cost
                                current_balance = user_balance - total_task_cost
                                current_balance = round(current_balance, 2)
                                record.tamount = current_balance # Update user balance in database
                                db_session.commit()

                            #the_raven(version=digit,
                            #    model_id="Hausa",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
                            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = os.path.join('static/audios',f'madugu_{digit}.wav')
                            return render_template('inference.html', greetings=greetings, balance=tamount, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                            #elif langtag != "" and langtag == "sw":
                            #    audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            #    create_speech_from_text(sentences, langtag, langiso, audio_path)  
                            #the_raven(version=digit,
                            #    model_id="Swahili",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
               	            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = os.path.join('static/audios',f'tai_{digit}.wav')
                            #    return render_template('inference.html', greetings=greetings, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                    elif is_matching_doc is True:
	                    #print('Request: '+ str(is_matching_doc)+str(langtag))
                        digit = random.randint(0, 1000)
                        exec_device = "cuda" if torch.cuda.is_available() else "cpu"
                        message = "Your audio is ready. Listen to speech"
        
                        sentences = read_docfile(file_path)
                        if sentences is not None and langtag is not None and langiso is not None:
                            #if langtag != "" and langtag == "ha":
                            audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            create_speech_from_text(sentences, langtag, langiso, audio_path)

                            file_size = get_file_size(audio_path) # Get the file size in bytes
                            if file_size:
                                print(f"File size: {file_size} bytes. File size is valid.")
                                task_cost_per_mb = 1.5  # Cost per MB in credits
                                file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                                total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost
                                
                                # Modify and substract user's total amount for usage cost
                                current_balance = user_balance - total_task_cost
                                current_balance = round(current_balance, 2)
                                record.tamount = current_balance # Update user balance in database
                                db_session.commit()

                            #the_raven(version=digit,
                            #    model_id="Hausa",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
                            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = os.path.join('static/audios',f'madugu_{digit}.wav')
                            return render_template('inference.html', greetings=greetings, balance=tamount, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                            #elif langtag != "" and langtag == "sw":
                            #    audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            #    create_speech_from_text(sentences, langtag, langiso, audio_path)  
                            #the_raven(version=digit,
                            #    model_id="Swahili",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
               	            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = os.path.join('static/audios',f'tai_{digit}.wav')
                            #    return render_template('inference.html', greetings=greetings, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                    elif is_matching_docx is True:
	                    #print('Request: '+ str(is_matching_docx)+str(langtag))	
                        digit = random.randint(0, 1000)
                        exec_device = "cuda" if torch.cuda.is_available() else "cpu"
                        message = "Your audio is ready. Listen to speech"
        
                        sentences = read_docxfile(file_path)
                        if sentences is not None and langtag is not None and langiso is not None:
                            #if langtag != "" and langtag == "ha":
                            audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            create_speech_from_text(sentences, langtag, langiso, audio_path)

                            file_size = get_file_size(audio_path) # Get the file size in bytes
                            if file_size:
                                print(f"File size: {file_size} bytes. File size is valid.")
                                task_cost_per_mb = 1.5  # Cost per MB in credits
                                file_size_mb = file_size / 1024 / 1024  # Get the file size in MB
                                total_task_cost = task_cost_per_mb * file_size_mb # Calculate the task cost

                                # Modify and substract user's total amount for usage cost
                                current_balance = user_balance - total_task_cost
                                current_balance = round(current_balance, 2)
                                record.tamount = current_balance # Update user balance in database
                                db_session.commit()

                            #the_raven(version=digit,
                            #    model_id="Hausa",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
                            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = os.path.join('static/audios',f'madugu_{digit}.wav')
                            return render_template('inference.html', greetings=greetings, balance=tamount, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
                            #elif langtag != "" and langtag == "sw":
                            #    audio_path = os.path.join('static/audios',f'{langiso}_{str(uuid.uuid4())}.wav')
                            #    create_speech_from_text(sentences, langtag, langiso, audio_path)  
                            #the_raven(version=digit,
                            #    model_id="Swahili",
                            #    exec_device=exec_device,
                            #    speed_over_quality=exec_device != "cuda",
               	            #    exsentence=sentences,
                            #    langcode=langtag)
                            #audio_path = os.path.join('static/audios',f'tai_{digit}.wav')
                            #    return render_template('inference.html', greetings=greetings, langtag = langtag, audio_visible=True, audio_file=audio_path, message=message)
        except NoResultFound:
            print("No records found in the database.")
        except Exception as e:
            db_session.rollback()
            print(f"An error occurred: {e}")
        finally:
            db_session.close()

        #os.rmdir('static/audios')
        #os.rmdir('static/uploads/docs')
        return render_template('inference.html', greetings=greetings, balance=tamount, langtag = langtag, audio_visible=False, audio_file='', message='')
    else:
        #return "Not logged in | <a href='/login'>Login</a>"
        return redirect(url_for('login'))


# Starting Media Dubbbing Processing
def preprocess_char(text, lang=None):
    """
    Special treatement of characters in certain languages
    """
    print(lang)
    if lang == 'ron':
        text = text.replace("ț", "ţ")
    return text

class TextMapper(object):
    def __init__(self, vocab_file):
        self.symbols = [x.replace("\n", "") for x in open(vocab_file, encoding="utf-8").readlines()]
        self.SPACE_ID = self.symbols.index(" ")
        self._symbol_to_id = {s: i for i, s in enumerate(self.symbols)}
        self._id_to_symbol = {i: s for i, s in enumerate(self.symbols)}

    def text_to_sequence(self, text, cleaner_names):
        '''Converts a string of text to a sequence of IDs corresponding to the symbols in the text.
        Args:
        text: string to convert to a sequence
        cleaner_names: names of the cleaner functions to run the text through
        Returns:
        List of integers corresponding to the symbols in the text
        '''
        sequence = []
        clean_text = text.strip()
        for symbol in clean_text:
            symbol_id = self._symbol_to_id[symbol]
            sequence += [symbol_id]
        return sequence

    def uromanize(self, text, uroman_pl):
        iso = "xxx"
        with tempfile.NamedTemporaryFile() as tf, \
             tempfile.NamedTemporaryFile() as tf2:
            with open(tf.name, "w") as f:
                f.write("\n".join([text]))
            cmd = f"perl " + uroman_pl
            cmd += f" -l {iso} "
            cmd +=  f" < {tf.name} > {tf2.name}"
            os.system(cmd)
            outtexts = []
            with open(tf2.name) as f:
                for line in f:
                    line =  re.sub(r"\s+", " ", line).strip()
                    outtexts.append(line)
            outtext = outtexts[0]
        return outtext

    def get_text(self, text, hps):
        text_norm = self.text_to_sequence(text, hps.data.text_cleaners)
        if hps.data.add_blank:
            text_norm = commons.intersperse(text_norm, 0)
        text_norm = torch.LongTensor(text_norm)
        return text_norm

    def filter_oov(self, text):
        val_chars = self._symbol_to_id
        txt_filt = "".join(list(filter(lambda x: x in val_chars, text)))
        print(f"text after filtering OOV: {txt_filt}")
        return txt_filt

def preprocess_text(txt, text_mapper, hps, uroman_dir=None, lang=None):
    txt = preprocess_char(txt, lang=lang)
    is_uroman = hps.data.training_files.split('.')[-1] == 'uroman'
    if is_uroman:
        with tempfile.TemporaryDirectory() as tmp_dir:
            if uroman_dir is None:
                cmd = f"git clone git@github.com:isi-nlp/uroman.git {tmp_dir}"
                print(cmd)
                subprocess.check_output(cmd, shell=True)
                uroman_dir = tmp_dir
            uroman_pl = os.path.join(uroman_dir, "bin", "uroman.pl")
            print(f"uromanize")
            txt = text_mapper.uromanize(txt, uroman_pl)
            print(f"uroman text: {txt}")
    txt = txt.lower()
    txt = text_mapper.filter_oov(txt)
    return txt


def preprocess_speech_text(txt, text_mapper, hps, uroman_dir=None, lang=None):
    txt = preprocess_char(txt, lang=lang)
    is_uroman = hps.data.training_files.split('.')[-1] == 'uroman'
    if is_uroman:
        with tempfile.TemporaryDirectory() as tmp_dir:
            if uroman_dir is None:
                cmd = f"git clone git@github.com:isi-nlp/uroman.git {tmp_dir}"
                print(cmd)
                subprocess.check_output(cmd, shell=True)
                uroman_dir = tmp_dir
            uroman_pl = os.path.join(uroman_dir, "bin", "uroman.pl")
            print(f"uromanize")
            txt = text_mapper.uromanize(txt, uroman_pl)
            print(f"uroman text: {txt}")
    #txt = txt.lower()
    #txt = text_mapper.filter_oov(txt)
    return txt


@retry(tries=3, delay=2, backoff=2)
# Method for downloading speech models
def download(lang, tgt_dir="./"):
    lang_fn, lang_dir = os.path.join(tgt_dir, lang+'.tar.gz'), os.path.join(tgt_dir, lang)
    if not os.path.exists(lang_dir):
        cmd = ";".join([
            f"wget https://dl.fbaipublicfiles.com/mms/tts/{lang}.tar.gz -O {lang_fn}",
            f"tar zxvf {lang_fn}"
        ])
        print(f"Download model for language: {lang}")
        subprocess.check_output(cmd, shell=True)
        print(f"Model checkpoints in {lang_dir}: {os.listdir(lang_dir)}")
        return lang_dir
    else:
        return lang_dir

#LANG = "hau"
#ckpt_dir = download(LANG)


# Method for extracting audio from video
def extract_audio_from_video(video_file):
    try:
        print("Extracting audio track")
        video = VideoFileClip(video_file)
        audio = video.audio
        audio_file = os.path.splitext(video_file)[0] + ".wav"
        audio.write_audiofile(audio_file)
        return audio_file
    except Exception as e:
        print(f"Error extracting audio from video: {e}")
        return None

@retry(tries=3, delay=2, backoff=2)
# Method for transcribing extracted audio using whisper
def transcribe_audio(audio_file, source_language):
    try:
        print("Transcribing audio track")
        model = whisper.load_model("small")
        trans = model.transcribe(audio_file, language=source_language, verbose=False, word_timestamps=True)
        return trans
    except Exception as e:
        print(f"Error transcribing audio: {e}")
        return None


@retry(tries=3, delay=2, backoff=2)
# Method for translating transcribed text
def translate_text(chunk, source_language, target_language):
    try:
        #translate_client = translate.Client()
        #results = translate_client.translate(texts, target_language=target_language)
        #return [result['translatedText'] for result in results]
        translator = Translator(from_lang=source_language, to_lang=target_language)
        translated_chunk = [translator.translate(sentence) for sentence in chunk]
        return [t for t in translated_chunk]
    except Exception as e:
        print(f"Error translating texts: {e}")
        return None



# Method for creating speech from text
def create_speech_from_text(text, target_language, isocode, audio_file):
    # Initialize an empty array to store audio segments
    audio_segments = []
    try:
        ckpt_dir = download(isocode)

        if torch.cuda.is_available():
          device = torch.device("cuda")
        else:
          device = torch.device("cpu")

        print(f"Run inference with {device}")
        vocab_file = f"{ckpt_dir}/vocab.txt"
        config_file = f"{ckpt_dir}/config.json"
        assert os.path.isfile(config_file), f"{config_file} doesn't exist"
        hps = utils.get_hparams_from_file(config_file)
        text_mapper = TextMapper(vocab_file)
        net_g = SynthesizerTrn(
        len(text_mapper.symbols),
        hps.data.filter_length // 2 + 1,
        hps.train.segment_size // hps.data.hop_length,**hps.model)
        net_g.to(device)
        _ = net_g.eval()

        g_pth = f"{ckpt_dir}/G_100000.pth"
        print(f"load {g_pth}")

        _ = utils.load_checkpoint(g_pth, net_g, None)

        print(f"text: {text}")
        audio_segments.clear()
        sentences = preprocess_speech_text(text, text_mapper, hps, lang=target_language)
        # Iterate through each sentence and generate audio
        for sentence in sentences:
            sentence_case = sentence.lower()
            text = text_mapper.filter_oov(sentence_case)
            stn_tst = text_mapper.get_text(text, hps)
            with torch.no_grad():
                x_tst = stn_tst.unsqueeze(0).to(device)
                x_tst_lengths = torch.LongTensor([stn_tst.size(0)]).to(device)
                hyp = net_g.infer(
                x_tst, x_tst_lengths, noise_scale=.667,
                noise_scale_w=0.8, length_scale=1.0
                )[0][0,0].cpu().float().numpy()

            # Append the generated audio to the list of audio segments
            audio_segments.append(hyp)
        
        # Concatenate the audio segments to create a single audio
        concatenated_audio = np.concatenate(audio_segments)
        # Save the generated audio as a WAV file
        wavsaver.write(audio_file, hps.data.sampling_rate, concatenated_audio)
        print(f"Generated audio saved {audio_file}")

        #Audio(hyp, rate=hps.data.sampling_rate)
        return audio_file
    except Exception as e:
        if os.path.isfile(audio_file):
            os.remove(audio_file)
        raise Exception(f"Error creating audio from text: {e}")



# Method for creating audio from text
def create_audio_from_text(text, target_language, isocode):
    audio_file = "translated_" + str(uuid.uuid4()) + ".wav"
    try:
        ckpt_dir = download(isocode)

        if torch.cuda.is_available():
          device = torch.device("cuda")
        else:
          device = torch.device("cpu")

        print(f"Run inference with {device}")
        vocab_file = f"{ckpt_dir}/vocab.txt"
        config_file = f"{ckpt_dir}/config.json"
        assert os.path.isfile(config_file), f"{config_file} doesn't exist"
        hps = utils.get_hparams_from_file(config_file)
        text_mapper = TextMapper(vocab_file)
        net_g = SynthesizerTrn(
        len(text_mapper.symbols),
        hps.data.filter_length // 2 + 1,
        hps.train.segment_size // hps.data.hop_length,**hps.model)
        net_g.to(device)
        _ = net_g.eval()

        g_pth = f"{ckpt_dir}/G_100000.pth"
        print(f"load {g_pth}")

        _ = utils.load_checkpoint(g_pth, net_g, None)

        print(f"text: {text}")
        txt = preprocess_text(text, text_mapper, hps, lang=target_language)
        stn_tst = text_mapper.get_text(txt, hps)
        with torch.no_grad():
          x_tst = stn_tst.unsqueeze(0).to(device)
          x_tst_lengths = torch.LongTensor([stn_tst.size(0)]).to(device)
          hyp = net_g.infer(
          x_tst, x_tst_lengths, noise_scale=.667,
          noise_scale_w=0.8, length_scale=1.0
          )[0][0,0].cpu().float().numpy()


        # Save the generated audio as a WAV file
        wavsaver.write(audio_file, hps.data.sampling_rate, hyp)
        print(f"Generated audio saved {audio_file}")

        #Audio(hyp, rate=hps.data.sampling_rate)

        return audio_file
    except Exception as e:
        if os.path.isfile(audio_file):
            os.remove(audio_file)
        raise Exception(f"Error creating audio from text: {e}")


ABBREVIATIONS = {
    "Mr.": "Mister",
    "Mrs.": "Misses",
    "No.": "Number",
    "Dr.": "Doctor",
    "Ms.": "Miss",
    "Ave.": "Avenue",
    "Blvd.": "Boulevard",
    "Ln.": "Lane",
    "Rd.": "Road",
    "a.m.": "before noon",
    "p.m.": "after noon",
    "ft.": "feet",
    "hr.": "hour",
    "min.": "minute",
    "sq.": "square",
    "St.": "street",
    "Asst.": "assistant",
    "Corp.": "corporation"
}

ISWORD = re.compile(r'.*\w.*')


# Method for retrieving sentences
def get_transcript_chunks(transcription, source_language):
    temp_files = []
    try:
        #ducked_audio = AudioSegment.from_wav(audio_file)
        if spacy_models[source_language] not in spacy.util.get_installed_models():
            spacy.cli.download(spacy_models[source_language])
        nlp = spacy.load(spacy_models[source_language])
        nlp.add_pipe("syllables", after="tagger")
        #merged_audio = AudioSegment.silent(duration=0)
        sentences = []
        sentence_starts = []
        sentence_ends = []
        sentence = ""
        sent_start = 0
        print("Composing sentences")
        for segment in tqdm(transcription["segments"]):
            if segment["text"].isupper():
                continue
            for i, word in enumerate(segment["words"]):
                if not ISWORD.search(word["word"]):
                    continue
                word["word"] = ABBREVIATIONS.get(word["word"].strip(), word["word"])
                if word["word"].startswith("-"):
                    sentence = sentence[:-1] + word["word"] + " "
                else:
                    sentence += word["word"] + " "
                # this is a trick to compensate the absense of VAD in Whisper
                word_syllables = sum(token._.syllables_count for token in nlp(word["word"]) if token._.syllables_count)
                segment_syllables = sum(token._.syllables_count for token in nlp(segment["text"]) if token._.syllables_count)
                if i == 0 or sent_start == 0:
                    word_speed = word_syllables / (word["end"] - word["start"])
                    if word_speed < 3:
                        sent_start = word["end"] - word_syllables / 3
                    else:
                        sent_start = word["start"]
                if i == len(segment["words"]) - 1:  # last word in segment
                    word_speed = word_syllables / (word["end"] - word["start"])
                    segment_speed = segment_syllables / (segment["end"] - segment["start"])
                    if word_speed < 1.0 or segment_speed < 2.0:
                        word["word"] += "."

                if word["word"].endswith("."):
                    sentences.append(sentence)
                    sentence_starts.append(sent_start)
                    sentence_ends.append(word["end"])
                    sent_start = 0
                    sentence = ""
        # retrieve sentences in chunks of 128
        print("Retrieving sentences")
        translated_texts = []
        for i in tqdm(range(0, len(sentences), 128)):
            chunk = sentences[i:i + 128]
            translated_texts.extend(chunk)

        return translated_texts
    except Exception as e:
        print(f"Error merging audio files: {e}")
        return None
    finally:
        # cleanup: remove all temporary files
        for file in temp_files:
            try:
                os.remove(file)
            except Exception as e:
                print(f"Error removing temporary file {file}: {e}")



# Method for merging and ducking audio files
def human_merge_audio(transcription, translated_texts, source_language, target_language, isocode, audio_file):
    temp_files = []
    try:
        ducked_audio = AudioSegment.from_wav(audio_file)
        if spacy_models[source_language] not in spacy.util.get_installed_models():
            spacy.cli.download(spacy_models[source_language])
        nlp = spacy.load(spacy_models[source_language])
        nlp.add_pipe("syllables", after="tagger")
        merged_audio = AudioSegment.silent(duration=0)
        sentences = []
        sentence_starts = []
        sentence_ends = []
        sentence = ""
        sent_start = 0
        print("Composing sentences")
        for segment in tqdm(transcription["segments"]):
            if segment["text"].isupper():
                continue
            for i, word in enumerate(segment["words"]):
                if not ISWORD.search(word["word"]):
                    continue
                word["word"] = ABBREVIATIONS.get(word["word"].strip(), word["word"])
                if word["word"].startswith("-"):
                    sentence = sentence[:-1] + word["word"] + " "
                else:
                    sentence += word["word"] + " "
                # this is a trick to compensate the absense of VAD in Whisper
                word_syllables = sum(token._.syllables_count for token in nlp(word["word"]) if token._.syllables_count)
                segment_syllables = sum(token._.syllables_count for token in nlp(segment["text"]) if token._.syllables_count)
                if i == 0 or sent_start == 0:
                    word_speed = word_syllables / (word["end"] - word["start"])
                    if word_speed < 3:
                        sent_start = word["end"] - word_syllables / 3
                    else:
                        sent_start = word["start"]
                if i == len(segment["words"]) - 1:  # last word in segment
                    word_speed = word_syllables / (word["end"] - word["start"])
                    segment_speed = segment_syllables / (segment["end"] - segment["start"])
                    if word_speed < 1.0 or segment_speed < 2.0:
                        word["word"] += "."

                if word["word"].endswith("."):
                    sentences.append(sentence)
                    sentence_starts.append(sent_start)
                    sentence_ends.append(word["end"])
                    sent_start = 0
                    sentence = ""
        print("Creating translated audio track")
        total_duration = 0
        prev_end_time = 0
        for i, translated_text in enumerate(tqdm(translated_texts)):
            translated_audio_file = create_audio_from_text(translated_text, target_language, isocode)
            #print(translated_audio_file)
            if translated_audio_file is None:
                raise Exception("Audio creation failed")
            temp_files.append(translated_audio_file)
            translated_audio = AudioSegment.from_wav(translated_audio_file)

            # Apply "ducking" effect: reduce volume of original audio during translated sentence
            start_time = int(sentence_starts[i] * 1000)
            end_time = start_time + len(translated_audio)
            
            #next_start_time = int(sentence_starts[i+1] * 1000) if i > len(translated_texts) - 1 else len(ducked_audio)
            #ducked_segment = ducked_audio[start_time:end_time].apply_gain(-10)  # adjust volume reduction as needed

            #fade_out_duration = min(500, max(1, start_time - prev_end_time))
            #fade_in_duration = min(500, max(1, next_start_time  - end_time))
            #prev_end_time = end_time
            
            # Apply fade in effect to the end of the audio before the ducked segment
            #if start_time == 0:
            #    ducked_audio = ducked_segment +  ducked_audio[end_time:].fade_in(fade_in_duration)
            #elif end_time == len(ducked_audio):
            #    ducked_audio = ducked_audio[:start_time].fade_out(fade_out_duration) + ducked_segment
            #else:
            #    ducked_audio = ducked_audio[:start_time].fade_out(fade_out_duration) \
            #                   + ducked_segment +  ducked_audio[end_time:].fade_in(fade_in_duration)            

            # Overlay the translated audio on top of the original audio
            ducked_audio = ducked_audio.overlay(translated_audio, position=total_duration)
            # Update the total duration with the duration of the newly added translated audio
            total_duration += len(translated_audio)
            
            original_duration = int(sentence_ends[i] * 1000)
            new_duration = len(translated_audio) + len(merged_audio)
            padding_duration = max(0, original_duration - new_duration)
            padding = AudioSegment.silent(duration=padding_duration)
            merged_audio += padding + translated_audio

        return merged_audio, ducked_audio
    except Exception as e:
        print(f"Error merging audio files (index {i}): {e}")
        return None
    finally:
        # cleanup: remove all temporary files
        for file in temp_files:
            try:
                os.remove(file)
            except Exception as e:
                print(f"Error removing temporary file {file}: {e}")


# Method for merging and ducking audio files
def merge_audio_files(transcription, source_language, from_lang, target_language, isocode, audio_file):
    temp_files = []
    try:
        ducked_audio = AudioSegment.from_wav(audio_file)
        if spacy_models[source_language] not in spacy.util.get_installed_models():
            spacy.cli.download(spacy_models[source_language])
        nlp = spacy.load(spacy_models[source_language])
        nlp.add_pipe("syllables", after="tagger")
        merged_audio = AudioSegment.silent(duration=0)
        sentences = []
        sentence_starts = []
        sentence_ends = []
        sentence = ""
        sent_start = 0
        print("Composing sentences")
        for segment in tqdm(transcription["segments"]):
            if segment["text"].isupper():
                continue
            for i, word in enumerate(segment["words"]):
                if not ISWORD.search(word["word"]):
                    continue
                word["word"] = ABBREVIATIONS.get(word["word"].strip(), word["word"])
                if word["word"].startswith("-"):
                    sentence = sentence[:-1] + word["word"] + " "
                else:
                    sentence += word["word"] + " "
                # this is a trick to compensate the absense of VAD in Whisper
                word_syllables = sum(token._.syllables_count for token in nlp(word["word"]) if token._.syllables_count)
                segment_syllables = sum(token._.syllables_count for token in nlp(segment["text"]) if token._.syllables_count)
                if i == 0 or sent_start == 0:
                    word_speed = word_syllables / (word["end"] - word["start"])
                    if word_speed < 3:
                        sent_start = word["end"] - word_syllables / 3
                    else:
                        sent_start = word["start"]
                if i == len(segment["words"]) - 1:  # last word in segment
                    word_speed = word_syllables / (word["end"] - word["start"])
                    segment_speed = segment_syllables / (segment["end"] - segment["start"])
                    if word_speed < 1.0 or segment_speed < 2.0:
                        word["word"] += "."

                if word["word"].endswith("."):
                    sentences.append(sentence)
                    sentence_starts.append(sent_start)
                    sentence_ends.append(word["end"])
                    sent_start = 0
                    sentence = ""
        # translate sentences in chunks of 128
        print("Translating sentences")
        translated_texts = []
        for i in tqdm(range(0, len(sentences), 128)):
            chunk = sentences[i:i + 128]
            print("Chunk to be translated:", chunk)  # Print the chunk of sentences
            translated_chunk = translate_text(chunk, from_lang, target_language)
            if translated_chunk is None:
                raise Exception("Translation failed")
            translated_texts.extend(translated_chunk)
        print(translated_texts)
        print("Creating translated audio track")
        total_duration = 0
        prev_end_time = 0
        for i, translated_text in enumerate(tqdm(translated_texts)):
            translated_audio_file = create_audio_from_text(translated_text, target_language, isocode)
            if translated_audio_file is None:
                raise Exception("Audio creation failed")
            temp_files.append(translated_audio_file)
            translated_audio = AudioSegment.from_wav(translated_audio_file)

            # Apply "ducking" effect: reduce volume of original audio during translated sentence
            start_time = int(sentence_starts[i] * 1000)
            end_time = start_time + len(translated_audio)
            #next_start_time = int(sentence_starts[i+1] * 1000) if i < len(translated_texts) - 1 else len(ducked_audio)
            #ducked_segment = ducked_audio[start_time:end_time].apply_gain(-10)  # adjust volume reduction as needed

            #fade_out_duration = min(500, max(1, start_time - prev_end_time))
            #fade_in_duration = min(500, max(1, next_start_time  - end_time))
            #prev_end_time = end_time
            # Apply fade in effect to the end of the audio before the ducked segment
            #if start_time == 0:
            #    ducked_audio = ducked_segment +  ducked_audio[end_time:].fade_in(fade_in_duration)
            #elif end_time == len(ducked_audio):
            #    ducked_audio = ducked_audio[:start_time].fade_out(fade_out_duration) + ducked_segment
            #else:
            #    ducked_audio = ducked_audio[:start_time].fade_out(fade_out_duration) \
            #                   + ducked_segment +  ducked_audio[end_time:].fade_in(fade_in_duration)

            # Overlay the translated audio on top of the original audio
            ducked_audio = ducked_audio.overlay(translated_audio, position=total_duration)
            # Update the total duration with the duration of the newly added translated audio
            total_duration += len(translated_audio)

            original_duration = int(sentence_ends[i] * 1000)
            new_duration = len(translated_audio) + len(merged_audio)
            padding_duration = max(0, original_duration - new_duration)
            padding = AudioSegment.silent(duration=padding_duration)
            merged_audio += padding + translated_audio

        return merged_audio, ducked_audio
    except Exception as e:
        print(f"Error merging audio files: {e}")
        return None
    finally:
        # cleanup: remove all temporary files
        for file in temp_files:
            try:
                os.remove(file)
            except Exception as e:
                print(f"Error removing temporary file {file}: {e}")


def save_audio_to_file(audio, filename):
    try:
        audio.export(filename, format="wav")
        print(f"Audio track with translation only saved to {filename}")
    except Exception as e:
        print(f"Error saving audio to file: {e}")


def replace_audio_in_video(video_file, serial, new_audio):
    try:
        # Load the video
        video = VideoFileClip(video_file)

        # Save the new audio to a temporary file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as temp_audio_file:
            new_audio.export(temp_audio_file.name, format="wav")
        #new_audio.export(str(uuid.uuid4())+"_duckled.wav", format="wav")

        # Load the new audio into an AudioFileClip
        try:
            new_audio_clip = AudioFileClip(temp_audio_file.name)
        except Exception as e:
            print(f"Error loading new audio into an AudioFileClip: {e}")
            return

        # Check if the audio is compatible with the video
        if new_audio_clip.duration < video.duration:
            print("Warning: The new audio is shorter than the video. The remaining video will have no sound.")
        elif new_audio_clip.duration > video.duration:
            print("Warning: The new audio is longer than the video. The extra audio will be cut off.")
            new_audio_clip = new_audio_clip.subclip(0, video.duration)

        # Set the audio of the video to the new audio
        video = video.set_audio(new_audio_clip)

        # Write the result to a new video file
        output_filename = os.path.splitext(video_file)[0] + serial + "_translated.mp4"
        try:
            video.write_videofile(output_filename, audio_codec='aac')
        except Exception as e:
            print(f"Error writing the new video file: {e}")
            return

        print(f"Translated video saved as {output_filename}")

    except Exception as e:
        print(f"Error replacing audio in video: {e}")
    finally:
        # Remove the temporary audio file
        if os.path.isfile(temp_audio_file.name):
            os.remove(temp_audio_file.name)


if __name__ == '__main__':
    os.makedirs('temp', exist_ok=True)
    # Create the table in the database (if it doesn't exist)
    Base.metadata.create_all(engine)
    app.run()
